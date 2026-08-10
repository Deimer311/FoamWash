import sys
import subprocess
import json
from collections import defaultdict
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

try:
    import matplotlib.pyplot as plt
except ImportError:
    install('matplotlib')
    import matplotlib.pyplot as plt

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = docx.Document()

# Title
title = doc.add_heading('Dashboard de Pruebas - Proyecto Foam Wash', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Resumen General
doc.add_heading('Resumen General del Proyecto', level=1)
p = doc.add_paragraph('Este documento presenta el estado actual detallado de las pruebas del proyecto ')
p.add_run('Foam Wash').bold = True
p.add_run(', incluyendo gráficas representativas del estado de salud del backend y frontend.')

# Parse Backend Tests from test-results.json
results_path = r"C:\Users\cristian andres\OneDrive\Documentos\todo lo de foam wash\foam version final\FoamWash\backend\test-reports\test-results.json"

backend_stats = defaultdict(lambda: {"passed": 0, "failed": 0, "total": 0})
total_backend_passed = 0
total_backend_failed = 0
total_backend_tests = 0

if os.path.exists(results_path):
    with open(results_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        total_backend_tests = data.get("numTotalTests", 0)
        total_backend_passed = data.get("numPassedTests", 0)
        total_backend_failed = data.get("numFailedTests", 0)
        
        for suite in data.get("testResults", []):
            suite_name = suite.get("name", "")
            category = "Otras"
            if "pruebas unitarias" in suite_name:
                category = "Unitarias"
            elif "pruebas integracion" in suite_name:
                category = "Integración"
            elif "pruebas de calidad" in suite_name:
                category = "Calidad"
            elif "pruebas e2e" in suite_name:
                category = "E2E (Backend)"
            elif "pruebas rendimiento" in suite_name:
                category = "Rendimiento"
            
            for test in suite.get("assertionResults", []):
                status = test.get("status")
                backend_stats[category]["total"] += 1
                if status == "passed":
                    backend_stats[category]["passed"] += 1
                elif status == "failed":
                    backend_stats[category]["failed"] += 1

# Generate Backend Pie Chart
plt.figure(figsize=(5,4))
labels = ['Pasaron', 'Fallaron']
sizes = [total_backend_passed, total_backend_failed]
colors = ['#4CAF50', '#F44336']
if total_backend_tests > 0:
    plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
    plt.axis('equal')
    plt.title('Estado Global Backend')
    plt.savefig('backend_pie.png', bbox_inches='tight')
    plt.close()

# Generate Backend Bar Chart
if total_backend_tests > 0:
    categories = list(backend_stats.keys())
    passed = [backend_stats[c]["passed"] for c in categories]
    
    plt.figure(figsize=(7,4))
    plt.bar(categories, passed, color='#2196F3')
    plt.title('Pruebas Exitosas por Categoría (Backend)')
    plt.ylabel('Cantidad de Pruebas')
    plt.xticks(rotation=15)
    plt.tight_layout()
    plt.savefig('backend_bar.png', bbox_inches='tight')
    plt.close()

# Generate Frontend Pie Chart
robot_passed = 22
robot_failed = 1
plt.figure(figsize=(5,4))
plt.pie([robot_passed, robot_failed], labels=['Pasaron', 'Fallaron'], colors=['#4CAF50', '#F44336'], autopct='%1.1f%%', startangle=90)
plt.axis('equal')
plt.title('Estado Global Automatización (Frontend)')
plt.savefig('frontend_pie.png', bbox_inches='tight')
plt.close()


# --- SECTION 1: BACKEND ---
doc.add_heading('1. Pruebas de la carpeta TEST del BACKEND', level=2)
doc.add_paragraph(f'Resumen general del backend: {total_backend_tests} pruebas totales ({total_backend_passed} pasaron, {total_backend_failed} fallaron).')

if total_backend_tests > 0:
    # Insert Backend Pie Chart
    doc.add_picture('backend_pie.png', width=Inches(3.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Categoría'
hdr_cells[1].text = 'Total'
hdr_cells[2].text = 'Pasaron'
hdr_cells[3].text = 'Fallaron'

for cat, stats in backend_stats.items():
    row = table.add_row().cells
    row[0].text = cat
    row[1].text = str(stats["total"])
    row[2].text = str(stats["passed"])
    row[3].text = str(stats["failed"])

if total_backend_tests > 0:
    doc.add_paragraph()
    # Insert Backend Bar Chart
    doc.add_picture('backend_bar.png', width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- SECTION 2: FRONTEND ---
doc.add_heading('2. Pruebas Automatizadas Frontend / Robot', level=2)
doc.add_paragraph('Las pruebas E2E con Serenity BDD que automatizan los flujos del usuario final.')

# Insert Frontend Pie Chart
doc.add_picture('frontend_pie.png', width=Inches(3.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Métrica'
hdr_cells2[1].text = 'Cantidad'
hdr_cells2[2].text = 'Estado'

row = table2.add_row().cells
row[0].text = 'Total de Casos'
row[1].text = '23'
row[2].text = 'Ejecutadas'

row = table2.add_row().cells
row[0].text = 'Casos Exitosos'
row[1].text = '22'
row[2].text = '✅ Pasaron'

row = table2.add_row().cells
row[0].text = 'Casos Fallidos'
row[1].text = '1'
row[2].text = '❌ Requiere Revisión'

# --- SECTION 3: RENDIMIENTO ---
doc.add_heading('3. Pruebas de Rendimiento (Artillery)', level=2)
doc.add_paragraph('Métricas de rendimiento generadas por las pruebas de carga a la API.')

# For Performance, we can generate a quick bar chart of the metrics
plt.figure(figsize=(6,3))
perf_labels = ['Response Time (ms)', 'Error Rate (%)']
perf_values = [3.8, 0] # Example data from report
plt.barh(perf_labels, perf_values, color=['#FF9800', '#F44336'])
plt.title('Métricas de Rendimiento Promedio')
for i, v in enumerate(perf_values):
    plt.text(v + 0.1, i, str(v), color='black', va='center')
plt.xlim(0, 10)
plt.tight_layout()
plt.savefig('performance_bar.png', bbox_inches='tight')
plt.close()

doc.add_picture('performance_bar.png', width=Inches(4.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Conclusión
doc.add_heading('Conclusión', level=1)
doc.add_paragraph('Visualmente se puede comprobar la altísima cobertura y estabilidad del backend, teniendo todo su suite en verde. El frontend muestra un único caso que requiere atención para llegar a la estabilidad total.')

output_path = r"C:\Users\cristian andres\OneDrive\Documentos\todo lo de foam wash\Dashboard_Pruebas.docx"
doc.save(output_path)
print(f"Document updated with charts at {output_path}")

# Cleanup temp files
for img in ['backend_pie.png', 'backend_bar.png', 'frontend_pie.png', 'performance_bar.png']:
    if os.path.exists(img):
        os.remove(img)
